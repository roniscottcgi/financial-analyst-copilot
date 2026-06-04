import streamlit as st

from typing import Any
from docx import Document
from langchain_classic.chains.combine_documents import create_stuff_documents_chain
from langchain_classic.chains.retrieval import create_retrieval_chain
from langchain_core.documents import Document as LCDocument
from langchain_core.messages import SystemMessage, HumanMessage, AIMessage
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.tools import tool
from langchain_text_splitters import RecursiveCharacterTextSplitter
from streamlit.elements.widgets.chat import ChatInputValue
from streamlit.runtime.uploaded_file_manager import UploadedFile

from src.service.llm_service import LLMService
from src.utils.factory import init_openai_client, run_read_only_query
from docx import Document
from pypdf import PdfReader

from utils.factory import append_to_vector_store, get_vector_store


class ChatUI:
    def __init__(self, llm_service: LLMService):
        self.llm_service = llm_service
        self.vector_store = None

    def collect_assistant_response(self, message: list[dict[str, Any]]):
        return self.llm_service.send_request_to_assistant(
            model=st.session_state["openai_model"],
            message=message)

    @staticmethod
    def extract_text(uploaded_file):
        name = uploaded_file.name
        if name.endswith('.pdf'):
            pdf_reader = PdfReader(uploaded_file)
            return "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
        elif name.endswith('.docx'):
            doc = Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs])
        return ""

    @staticmethod
    def append_chat_messages(role: str, user_input: str | ChatInputValue | list[Any]):
        st.session_state.messages.append({"role": role, "content": user_input})

    @staticmethod
    def display_chat_history():
        for i, message in enumerate(st.session_state.messages):
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

    def render(self):
        st.title("Financial Assistant Chat")

        if "show_banner" not in st.session_state:
            st.session_state.show_banner = True
        if "messages" not in st.session_state:
            st.session_state.messages = []
        if "indexed_files" not in st.session_state:
            st.session_state.indexed_files = set()

        chat_toggle = self.toggle_container()

        if not chat_toggle:
            st.info("Chat is currently closed. Toggle the chat window to open the chat assistant.")
            return

        uploaded_files = None
        chat_container = st.container(height=500, autoscroll=True)
        new_user_input = st.chat_input("How can I assist you?", key="chat_input_key")
        uploaded_files = st.file_uploader(
            "Upload your data file(s)",
            type=["docx", "pdf"],
            accept_multiple_files=True)

        if uploaded_files:
            self.index_uploaded_files(uploaded_files)

        if self.vector_store is not None:
            # st.caption(":green[✔ Vector Database Status: Active & Loaded]")
            st.markdown(
                '<p style="color: #2ea44f; font-size: 0.85rem; margin: 0;">'
                '<span style="font-size: 1.5rem; vertical-align: middle; padding-right: 4px;">✔</span> '
                'Vector Database Status: Active & Loaded'
                '</p>',
                unsafe_allow_html=True)

        assistant_payload = None

        if new_user_input:
            self.append_chat_messages("user", new_user_input)
            assistant_payload = [
                {"role": m["role"], "content": m["content"]}
                for m in st.session_state.messages
            ]

        with chat_container:
            if st.session_state.show_banner and not self.vector_store:
                with st.container():
                    # Use columns to position an "Exit" or "X" button on the right
                    col1, col2 = st.columns([2, .5])
                    with col1:
                        st.info("💡Upload documents below for RAG-based results.")
                    with col2:
                        # When clicked, this button triggers a rerun and hides the banner
                        if st.button("✖️", key="exit_banner"):
                            st.session_state.show_banner = False
                            st.rerun()
            self.display_chat_history()

            if assistant_payload:
                with st.chat_message("assistant"):
                    # Fallback if user types without uploading docs
                    if uploaded_files is None:
                        stream = self.collect_assistant_response(assistant_payload)
                        response = st.write_stream(stream)
                        self.append_chat_messages("assistant", response)
                    else:
                        user_docs_store = get_vector_store("user_documents")
                        db_schema_store = get_vector_store("db_schema")

                        # 1. Fetch relevant background contexts
                        user_docs_results = user_docs_store.similarity_search_with_score(new_user_input, k=3)
                        db_schema_results = db_schema_store.similarity_search_with_score(new_user_input, k=4)

                        # DEBUG: Print exactly what chunks the model is seeing
                        print("--- CODESPACE DEBUG: SCHEMA CHUNKS RETRIEVED ---")
                        for doc, score in db_schema_results:
                            print(f"Score: {score} | Content: {doc.page_content}")
                        print("------------------------------------------------")

                        # Extract text from vector fragments to format into a clean schema reference
                        schema_context = "\n".join([doc.page_content for doc, _ in db_schema_results])
                        doc_context = "\n".join([doc.page_content for doc, _ in user_docs_results])

                        @tool
                        def run_database_query(sql_query: str) -> str:
                            """Executes a valid SQL SELECT query against the underlying database tables to extract live figures and rows."""
                            return run_read_only_query(sql_query)

                        tools = [run_database_query]

                        # 3. Get raw model client and bind tools capability
                        llm_client = init_openai_client()
                        # Bind tools ensures the model knows it can call 'run_database_query'
                        llm_with_tools = llm_client.bind_tools(tools)

                        # 4. Construct a dynamic system message containing instructions and schema references
                        system_prompt = (
                            "You are a helpful Financial Assistant with read-only access to a database schema and user documents.\n\n"
                            f"--- AVAILABLE DATABASE SCHEMA MATCHES ---\n{schema_context}\n\n"
                            f"--- USER DOCUMENT CONTENT ---\n{doc_context}\n\n"
                            "INSTRUCTIONS:\n"
                            "1. If answering the user's prompt requires extracting metrics, lists, or aggregates from tables, generate a valid SQL query "
                            "and invoke the 'run_database_query' tool.\n"
                            "2. Do not write example code or say you lack access. Use your tool to grab data.\n"
                            "3. If the document content contains the direct answer, answer from the documents.\n"
                            "4. Once tool execution data is returned to you, synthesize a clear response for the user."
                        )

                        # Re-map chat history to official LangChain message classes
                        from langchain_core.messages import SystemMessage, HumanMessage, AIMessage, ToolMessage

                        langchain_messages = [SystemMessage(content=system_prompt)]
                        for m in st.session_state.messages[:-1]:  # Append previous conversational history
                            if m["role"] == "user":
                                langchain_messages.append(HumanMessage(content=m["content"]))
                            else:
                                langchain_messages.append(AIMessage(content=m["content"]))
                        # Add current user prompt
                        langchain_messages.append(HumanMessage(content=new_user_input))

                        # 5. First Inference Pass: Let LLM determine intent and possibly make a tool call
                        placeholder = st.empty()
                        with st.spinner("Analyzing request..."):
                            ai_msg = llm_with_tools.invoke(langchain_messages)

                        # Check if the LLM chose to call our Python function
                        if ai_msg.tool_calls:
                            langchain_messages.append(ai_msg)  # Save the tool call request

                            for tool_call in ai_msg.tool_calls:
                                if tool_call["name"] == "run_database_query":
                                    query_to_run = tool_call["args"]["sql_query"]

                                    # Show a status banner in the Streamlit UI so the user knows what query is running
                                    with st.status(f"🤖 Agent executing SQL query...", expanded=False) as status:
                                        st.code(query_to_run, language="sql")
                                        # Trigger our backend database retrieval function
                                        live_db_data = run_database_query.invoke(tool_call)
                                        status.update(label="Query complete! Synthesizing data...", state="complete")

                                    # Append raw rows return packet back to model memory context
                                    langchain_messages.append(live_db_data)

                            # 6. Second Inference Pass: Stream the final answer back to the UI using the fresh rows data
                            full_response = ""
                            for chunk in llm_client.stream(langchain_messages):
                                if chunk.content:
                                    full_response += chunk.content
                                    placeholder.markdown(full_response)

                            self.append_chat_messages("assistant", full_response)

                        else:
                            # Standard Fallback: The model answered directly without needing live data tools
                            placeholder.markdown(ai_msg.content)
                            self.append_chat_messages("assistant", ai_msg.content)

    def index_uploaded_files(self, uploaded_files: list[UploadedFile] | UploadedFile):
        all_lc_docs = []
        chunks = []
        stable_ids = []

        for i, uploaded_file in enumerate(uploaded_files):
            if uploaded_file.name in st.session_state.indexed_files:
                continue

            text = self.extract_text(uploaded_file)
            if text.strip():
                doc = LCDocument(page_content=text, metadata={"source": uploaded_file.name})
                all_lc_docs.append(doc)
            if all_lc_docs:
                text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
                chunks = text_splitter.split_documents(all_lc_docs)
                doc_id = f"{i}:{uploaded_file.name}"
                stable_ids.append(doc_id)
            st.session_state.indexed_files.add(uploaded_file.name)

        with st.spinner("Adding files to store..."):
            self.vector_store = append_to_vector_store(
                chunks=chunks,
                ids=stable_ids,
                collection_name="user_documents")

            st.success(f"Successfully indexed {len(chunks)} text chunks into Chroma!")
            # st.rerun()

    def toggle_container(self) -> bool:
        with st.container(key="toggle_container"):
            col1, col2 = st.columns([3, 1])

            with col1:
                chat_toggle = st.toggle("Show Chat Assistant", key="sidebar_toggle")

            with col2:
                if chat_toggle and st.session_state.messages:
                    if st.button("Clear", key="clear_chat_btn", use_container_width=True):
                        st.session_state.messages = []
                        st.rerun()
        return chat_toggle